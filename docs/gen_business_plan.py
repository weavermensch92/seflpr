"""
SelfPR 사업계획서 DOCX 생성 스크립트
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── 페이지 여백 설정 ───────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

# ─── 스타일 헬퍼 ────────────────────────────────────────────
BLUE   = RGBColor(0x1E, 0x40, 0xAF)   # #1E40AF 진한 파랑
SLATE  = RGBColor(0x47, 0x55, 0x69)   # 슬레이트 회색
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xEF, 0xF6, 0xFF)   # 연한 배경
GREEN  = RGBColor(0x16, 0xA3, 0x4A)
RED    = RGBColor(0xDC, 0x26, 0x26)
AMBER  = RGBColor(0xD9, 0x77, 0x06)

def add_heading(text, level=1, color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    if level == 1:
        run.font.size = Pt(16)
        # 하단 테두리
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '1E40AF')
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = SLATE
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = SLATE

def add_para(text, bold=False, color=None, indent=False, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.7)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color

def add_bullet(text, level=0, color=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = color

def shade_row(row, hex_color):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

def make_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    hrow = table.rows[0]
    shade_row(hrow, '1E40AF')
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = WHITE
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 데이터 행
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        if ri % 2 == 0:
            shade_row(row, 'EFF6FF')
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 열 너비
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # 간격
    return table


# ═══════════════════════════════════════════════════════════
# 표지
# ═══════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.space_after = Pt(6)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SelfPR")
run.bold = True
run.font.size = Pt(40)
run.font.color.rgb = BLUE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("AI 기반 자기소개서 자동 작성 & 면접 코칭 플랫폼")
r2.font.size = Pt(16)
r2.font.color.rgb = SLATE

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("사  업  계  획  서")
r3.bold = True
r3.font.size = Pt(22)
r3.font.color.rgb = BLUE

doc.add_paragraph()
doc.add_paragraph()

for line in ["버전: v1.0", "작성일: 2026년 3월", "분류: AI SaaS / HR Tech / EdTech"]:
    pp = doc.add_paragraph()
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(line)
    rr.font.size = Pt(11)
    rr.font.color.rgb = SLATE

doc.add_page_break()


# ═══════════════════════════════════════════════════════════
# 1. 서비스 요약
# ═══════════════════════════════════════════════════════════
add_heading("1. 서비스 요약 (Executive Summary)")

add_para(
    "SelfPR은 한국 취업 준비생을 위한 AI 기반 자기소개서 자동 작성 + 면접 예상 질문 연습 플랫폼입니다.",
    bold=True
)
add_para(
    "기업명과 지원 포지션을 입력하면 AI가 해당 기업을 직접 리서치하고, 유저의 학력·경력·프로젝트 "
    "데이터를 조합하여 기업 방향에 맞는 스토리텔링 자기소개서를 자동 생성합니다. 생성된 자소서를 "
    "기반으로 HR + 기술 담당자 이중 시각의 면접 예상 질문을 출제하고, 유저의 답변에 대해 실시간 "
    "피드백을 제공하는 1문 1답 연습 기능까지 제공합니다."
)
add_para(
    "핵심 가치: 단순 글쓰기 도구가 아닌, 기업을 스스로 조사하고 판단하는 능동적 AI 취업 코치.",
    bold=True, color=BLUE
)

doc.add_paragraph()


# ═══════════════════════════════════════════════════════════
# 2. 문제 정의
# ═══════════════════════════════════════════════════════════
add_heading("2. 문제 정의")

add_heading("2-1. 현재 취업 준비의 문제점", level=2)

add_para("■ 자기소개서 작성", bold=True, color=BLUE)
for b in [
    "취업 준비생 1인당 평균 12개 기업에 지원 (2025 취업포털 설문)",
    "기업마다 다른 인재상, 포지션별 다른 요구사항 → 각 자소서에 맞춤화 필요",
    "맞춤화된 자소서 1개 작성에 평균 4-8시간 소요",
    "ChatGPT로 작성 시 AI 어투가 그대로 드러나 감점 요인",
]:
    add_bullet(b)

add_para("■ 면접 준비", bold=True, color=BLUE)
for b in [
    "자소서 기반의 맞춤형 예상 질문 없음 → 일반적인 면접 질문집에 의존",
    "혼자 연습할 때 피드백 수단 없음",
    "꼬리 질문 대응 훈련 어려움 → 실전에서 당황",
]:
    add_bullet(b)

add_para("■ 개인정보 보안", bold=True, color=BLUE)
for b in [
    "기존 자소서 작성 도구는 학력·경력 등 민감 정보를 서버에 무단 저장",
    "유저는 자신의 개인정보가 어디에 저장되는지 모름",
]:
    add_bullet(b)

add_heading("2-2. 시장 규모", level=2)
make_table(
    ["구분", "수치", "출처"],
    [
        ["연간 취업 지원 건수 (한국)", "약 800만 건", "통계청 2025"],
        ["취업 준비생 수", "약 320만 명", "통계청 2025"],
        ["HR Tech 국내 시장", "약 2조 원 (2025)", "KDB미래전략연구소"],
        ["AI 글쓰기 도구 글로벌 시장", "약 2.5조 원 (2025)", "Grand View Research"],
        ["연 성장률", "27% CAGR", "Grand View Research"],
    ],
    col_widths=[6, 5, 5]
)

add_heading("2-3. 타깃 유저", level=2)
add_para("Primary: 대학교 3-4학년 및 졸업 후 1-2년차 취업 준비생", bold=True)
for b in ["연령대: 22-27세", "특징: 디지털 네이티브, AI 도구 친숙, 취업 압박 높음", "규모: 약 150만 명 (활성 취업 준비층)"]:
    add_bullet(b)

add_para("Secondary: 이직을 준비하는 직장인 (3-7년차)", bold=True)
for b in ["연령대: 27-35세", "특징: 포지션별 맞춤 자소서 필요, 시간 부족", "규모: 약 50만 명"]:
    add_bullet(b)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════
# 3. 솔루션
# ═══════════════════════════════════════════════════════════
add_heading("3. 솔루션")

add_heading("3-1. AI 자기소개서 자동 생성", level=2)
for b in [
    "기업명 + 포지션 입력 → AI 기업 리서치 (뉴스/공식사이트/합격사례 자동 수집)",
    "유저 프로필과 기업 요구 매칭 → STAR 기법 스토리텔링 자소서 초안 생성",
    "AI 리뷰 + 어투 검토 + 갭 분석 → PDF/Word 내보내기",
]:
    add_bullet(b)

add_para("차별점", bold=True, color=BLUE)
for b in [
    "기업 리서치를 AI가 직접 수행 → 유저가 조사할 필요 없음",
    "자소서를 쓰는 AI가 20년차 HR + 10년차 실무자 관점으로 작성",
    "AI 어투 감지 → 사람이 쓴 것처럼 자동 재작성",
]:
    add_bullet(b, color=GREEN)

add_heading("3-2. 면접 예상 질문 연습", level=2)
for b in [
    "자소서 최신 버전 자동 스냅샷 → 기업 가치관 + 면접 레퍼런스 서치",
    "1문 1답 출제 (자소서 / 기업가치관 / 직무 / 상황 4가지 유형)",
    "유저 답변 → 실시간 AI 피드백 (스트리밍)",
    "꼬리 질문 or 신규 질문 유저 선택 → 세션 요약 + 취약 영역 분석 + PDF",
]:
    add_bullet(b)

add_para("차별점", bold=True, color=BLUE)
for b in [
    "자소서 내용 기반 맞춤형 질문 (일반 질문집과 다름)",
    "답변을 직접 주지 않고 '대응력을 기르는' 피드백 방식",
    "꼬리 질문 훈련 → 실전 면접 압박 상황 대비",
]:
    add_bullet(b, color=GREEN)

add_heading("3-3. 개인정보 보호 우선 설계", level=2)
for b in [
    "파일 원본 서버 미저장: 업로드 파일은 메모리에서만 처리 후 즉시 폐기",
    "미결제 시 로컬 저장: 포인트 충전 전까지 프로필은 브라우저 localStorage만",
    "투명한 고지: '이 데이터는 현재 이 브라우저에만 저장됩니다' 항상 표시",
]:
    add_bullet(b)

doc.add_paragraph()


# ═══════════════════════════════════════════════════════════
# 4. 비즈니스 모델
# ═══════════════════════════════════════════════════════════
add_heading("4. 비즈니스 모델")

add_heading("4-1. 포인트 과금 구조", level=2)
make_table(
    ["항목", "포인트", "원화 환산"],
    [
        ["포인트 충전", "100P / 10,000원", "1P = 100원"],
        ["자소서 프로젝트 생성", "30P", "3,000원"],
        ["면접 연습 세션 시작", "60P", "6,000원"],
        ["면접 — 신규 질문", "3P", "300원"],
        ["면접 — 꼬리 질문", "1P (질문당 최대 5개)", "100원"],
    ],
    col_widths=[7, 4, 5]
)

add_heading("4-2. AI 원가 분석", level=2)
add_para("자소서 생성 1건당 AI 비용", bold=True)
make_table(
    ["단계", "모델", "비용"],
    [
        ["기업 리서치", "GPT-4o-mini", "~$0.003"],
        ["기업 분석", "Claude Sonnet 4.6", "~$0.050"],
        ["프로필 매칭", "GPT-4o-mini", "~$0.003"],
        ["자소서 작성 (6문항)", "Claude Sonnet 4.6", "~$0.070"],
        ["합계", "", "~$0.13 (약 180원)"],
    ],
    col_widths=[5.5, 5.5, 5]
)
add_para("수익 = 3,000원 − 180원 = 2,820원/건 (마진율 94%)", bold=True, color=GREEN)

doc.add_paragraph()
add_para("면접 연습 세션당 AI 비용 (20문답 기준)", bold=True)
make_table(
    ["항목", "비용"],
    [
        ["시스템 프롬프트 + 첫 질문", "~$0.03"],
        ["답변 피드백 × 20", "~$0.40"],
        ["신규 질문 × 10", "~$0.15"],
        ["꼬리 질문 × 10", "~$0.10"],
        ["합계", "~$0.68 (약 940원)"],
    ],
    col_widths=[10, 6]
)
add_para("세션 수입 ≈ 10,000원 / AI 원가 940원 → 수익 9,060원 (마진율 91%)", bold=True, color=GREEN)

add_heading("4-3. 기업 리서치 캐시 효과", level=2)
for b in [
    "동일 기업+포지션 2번째 요청부터 캐시 재활용 → AI 원가 0",
    "인기 기업(삼성전자, SK 등)은 수십 명이 공유 → 건당 원가 급감",
    "캐시 7일 TTL → 최신성 유지",
]:
    add_bullet(b)

add_heading("4-4. 수익 예측", level=2)
add_para("보수적 시나리오 (서비스 론칭 6개월)", bold=True)
make_table(
    ["지표", "수치"],
    [
        ["MAU", "1,000명"],
        ["유료 전환율", "20%"],
        ["월 결제 유저", "200명"],
        ["1인당 월평균 소비", "자소서 2건 + 면접 1세션 = 120P ≈ 20,000원"],
        ["월 매출", "약 400만 원"],
        ["AI 원가 (마진율 95%)", "약 20만 원"],
        ["월 영업이익", "약 380만 원"],
    ],
    col_widths=[7, 9]
)

doc.add_paragraph()
add_para("성장 시나리오 (12개월)", bold=True)
make_table(
    ["지표", "수치"],
    [
        ["MAU", "10,000명"],
        ["유료 전환율", "25%"],
        ["월 결제 유저", "2,500명"],
        ["1인당 월평균 소비", "20,000원"],
        ["월 매출", "약 5,000만 원"],
        ["연 매출", "약 6억 원"],
    ],
    col_widths=[7, 9]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════
# 5. 경쟁 분석
# ═══════════════════════════════════════════════════════════
add_heading("5. 경쟁 분석")

add_heading("5-1. 주요 경쟁자", level=2)
make_table(
    ["서비스", "특징", "약점"],
    [
        ["자소설닷컴", "국내 1위 자소서 첨삭 커뮤니티", "AI 없음, 사람이 첨삭 (느림, 비쌈)"],
        ["스펙업", "자소서 작성 AI 도구", "기업 리서치 없음, 면접 기능 없음"],
        ["ChatGPT", "범용 AI", "기업 맞춤화 없음, 개인정보 보호 취약"],
        ["노션 AI", "글쓰기 보조", "취업 특화 기능 없음"],
        ["링커리어", "채용 플랫폼 + 합격 자소서 열람", "자동 생성 없음"],
    ],
    col_widths=[4, 6, 6]
)

add_heading("5-2. SelfPR의 차별화 포인트", level=2)
make_table(
    ["항목", "SelfPR", "경쟁자"],
    [
        ["AI 기업 리서치 자동화", "✅", "❌"],
        ["자소서 기반 맞춤 면접 질문", "✅", "❌"],
        ["꼬리 질문 훈련", "✅", "❌"],
        ["AI 어투 감지 & 인간화", "✅", "❌"],
        ["개인정보 로컬 우선 보호", "✅", "❌ (대부분 서버 저장)"],
        ["포인트 유연 과금", "✅", "❌ (월정액 또는 건당 고정가)"],
    ],
    col_widths=[7, 4, 5]
)

add_heading("5-3. 진입장벽", level=2)
for i, b in enumerate([
    "기업 리서치 캐시 데이터: 사용자가 늘수록 캐시 품질 향상 → 신규 경쟁자 따라오기 어려움",
    "면접 Q&A 데이터: 누적될수록 질문 품질 개선 (학습 루프)",
    "프롬프트 엔지니어링: HR 페르소나 프롬프트 축적 → 어드민에서 지속 개선",
    "네트워크 효과: 합격 유저 후기 → 신뢰도 → 신규 유저 유입",
], 1):
    add_bullet(f"{i}. {b}")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════
# 6. 마케팅 전략
# ═══════════════════════════════════════════════════════════
add_heading("6. 마케팅 전략")

add_heading("6-1. 채널별 전략", level=2)
make_table(
    ["채널", "전략", "목표"],
    [
        ["에브리타임", "취업 준비 커뮤니티 바이럴", "대학생 유입"],
        ["블라인드", "취준/이직 게시판 자연스러운 홍보", "직장인 유입"],
        ["유튜브", "'AI로 삼성전자 자소서 쓰기' 리뷰 콘텐츠", "SEO + 브랜딩"],
        ["인스타그램", "합격 자소서 비포/애프터", "MZ 유입"],
        ["SEO", "'삼성전자 자소서', '현대 면접 질문' 등 키워드", "유기 트래픽"],
    ],
    col_widths=[4, 7, 5]
)

add_heading("6-2. 성장 지표 (KPI)", level=2)
make_table(
    ["지표", "1개월", "3개월", "6개월", "12개월"],
    [
        ["가입자", "500", "3,000", "10,000", "50,000"],
        ["MAU", "200", "1,500", "5,000", "20,000"],
        ["유료 전환율", "15%", "20%", "25%", "30%"],
        ["월 매출", "90만원", "600만원", "2,500만원", "1.2억원"],
    ],
    col_widths=[4, 3, 3, 3, 3]
)

add_heading("6-3. 바이럴 루프", level=2)
add_para(
    "합격 성공 유저 → '에브리타임/블라인드 후기' → 신규 유저 유입 → "
    "신규 가입 시 무료 포인트(5P) 지급 → 자소서 생성 → 합격 성공 → 루프 반복"
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════
# 7. 기술 역량
# ═══════════════════════════════════════════════════════════
add_heading("7. 기술 역량")

add_heading("7-1. 기술 스택", level=2)
make_table(
    ["영역", "기술", "이유"],
    [
        ["Frontend", "React + TypeScript + Vite + shadcn/ui", "빠른 빌드, 커스터마이즈 용이"],
        ["Backend", "FastAPI (Python)", "AI 라이브러리 호환, async"],
        ["DB", "PostgreSQL", "안정적, AES-256 암호화"],
        ["AI", "Claude Sonnet 4.6 + GPT-4o-mini", "자소서/면접 생성, 파일 파싱"],
        ["인증", "JWT RS256", "Access 30분 + Refresh 7일"],
        ["결제", "토스페이먼츠", "국내 PG사, 카드/계좌이체"],
        ["스트리밍", "FastAPI SSE", "면접 AI 피드백 실시간 스트리밍"],
    ],
    col_widths=[3, 7, 6]
)

add_heading("7-2. 현재 구현 완료 기능 (2026.03)", level=2)
make_table(
    ["기능", "구현 상태"],
    [
        ["AI 자소서 생성 (Claude Sonnet 4.6)", "✅ 완료"],
        ["AI 리뷰 + 버전 비교", "✅ 완료"],
        ["갭 분석", "✅ 완료"],
        ["AI 어투 감지 & 인간화", "✅ 완료"],
        ["프로필 파일 파싱 (PDF/Excel/Word/OCR)", "✅ 완료"],
        ["포인트 시스템 (서비스 레이어)", "✅ 완료"],
        ["어드민 프롬프트 관리", "✅ 완료"],
        ["로컬 프로필 저장 (localStorage)", "✅ 완료"],
        ["면접 연습 기능 (Phase 7)", "🔲 개발 중"],
        ["결제/충전 페이지", "🔲 개발 중"],
    ],
    col_widths=[10, 6]
)

add_heading("7-3. 개발 로드맵", level=2)
make_table(
    ["Phase", "내용", "예상 완료"],
    [
        ["Phase 5", "DB 마이그레이션 수정 (포인트 컬럼)", "2주"],
        ["Phase 6", "결제 연동 (토스페이먼츠) + 충전 페이지", "2주"],
        ["Phase 7", "면접 연습 전체 구현 (에이전트 + API + 프론트)", "3주"],
        ["Phase 8", "기업 리서치 파이프라인 고도화", "2주"],
        ["Phase 9", "보안 강화 + Celery + 배포", "2주"],
        ["MVP 완성", "", "2026년 5월 예정"],
    ],
    col_widths=[3, 9, 4]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════
# 8. 팀 구성
# ═══════════════════════════════════════════════════════════
add_heading("8. 팀 구성")

add_heading("8-1. 현재 (1인 개발)", level=2)
make_table(
    ["역할", "담당 업무"],
    [["창업자 / Full-stack", "기획 + 백엔드 + 프론트엔드 + AI 프롬프트 설계"]],
    col_widths=[6, 10]
)

add_heading("8-2. 투자 유치 후 계획 (6개월 내)", level=2)
make_table(
    ["역할", "인원", "주요 업무"],
    [
        ["CTO / AI Engineer", "1명", "AI 파이프라인 고도화, 기업 리서치 품질"],
        ["Frontend Engineer", "1명", "UI/UX 개선, 면접 연습 화면"],
        ["Growth Marketer", "1명", "에브리타임/블라인드 바이럴, 콘텐츠"],
        ["CS / Operations", "1명", "고객 대응, 합격 사례 수집"],
    ],
    col_widths=[5, 3, 8]
)

doc.add_paragraph()


# ═══════════════════════════════════════════════════════════
# 9. 재무 계획
# ═══════════════════════════════════════════════════════════
add_heading("9. 재무 계획")

add_heading("9-1. 초기 비용 구조", level=2)
make_table(
    ["항목", "월 비용"],
    [
        ["AWS (ECS + RDS + CloudFront)", "약 30만원"],
        ["Anthropic API (Claude Sonnet)", "매출 연동, 약 5%"],
        ["OpenAI API (GPT-4o-mini)", "매출 연동, 약 1%"],
        ["토스페이먼츠 수수료", "매출의 1.1-3.3%"],
        ["기타 (도메인, 모니터링)", "약 10만원"],
        ["고정비 합계", "약 40만원/월"],
    ],
    col_widths=[10, 6]
)

add_heading("9-2. 손익분기점", level=2)
add_para(
    "월 고정비 40만원 / 평균 마진 5,000원 = 월 80건 유료 거래 = 월 유료 유저 40명 (1인 2건 기준)"
)
add_para("손익분기점이 매우 낮아 초기 자금 리스크 최소화", bold=True, color=GREEN)

add_heading("9-3. 자금 조달 계획", level=2)
make_table(
    ["단계", "시기", "금액", "방법"],
    [
        ["부트스트랩", "현재", "0원", "자기 자본 + 개발 완료"],
        ["Pre-Seed", "MVP 출시 후", "1-3억", "엔젤 / 엑셀러레이터"],
        ["Seed", "12개월 후", "5-10억", "VC"],
    ],
    col_widths=[4, 4, 4, 4]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════
# 10. 리스크 & 대응
# ═══════════════════════════════════════════════════════════
add_heading("10. 리스크 & 대응")
make_table(
    ["리스크", "영향도", "대응 방안"],
    [
        ["AI API 가격 인상", "중", "다중 모델 지원 (Claude/GPT 전환 가능 구조)"],
        ["경쟁자 진입 (대형 플랫폼)", "고", "기업 리서치 캐시 + 면접 데이터 선점 → 전환 비용 높임"],
        ["개인정보 규제 강화", "중", "로컬 저장 정책 + 원본 파일 미저장 구조 이미 구현"],
        ["AI 자소서 감지/블라인드 처리", "고", "AI 어투 제거 기능 핵심화, 사람이 쓴 것처럼 검증"],
        ["ChatGPT 고도화", "중", "기업 맞춤화 + 면접 연동 + 한국 HR 특화로 차별화"],
    ],
    col_widths=[5, 2, 9]
)

doc.add_paragraph()


# ═══════════════════════════════════════════════════════════
# 11. 성공 기준
# ═══════════════════════════════════════════════════════════
add_heading("11. 성공 기준")

make_table(
    ["기간", "목표"],
    [
        ["단기 (6개월)", "MAU 5,000 / 유료 전환 25% / 월 매출 2,500만원 / 합격 사례 100건"],
        ["중기 (12개월)", "MAU 20,000 / 월 매출 1.2억 / 기업 캐시 500개 / Seed 투자 유치"],
        ["장기 (3년)", "연 매출 50억 / B2B 확장 (대학·기업) / 일본 글로벌 진출"],
    ],
    col_widths=[4, 12]
)

doc.add_paragraph()


# ═══════════════════════════════════════════════════════════
# 12. 결론
# ═══════════════════════════════════════════════════════════
add_heading("12. 결론")

add_para(
    "SelfPR은 취업 준비생의 가장 큰 고통점 — 맞춤형 자소서 작성과 면접 준비 — 을 AI로 해결하는 서비스입니다.",
    bold=True
)
add_para(
    "단순한 ChatGPT 래퍼가 아닌, 기업을 스스로 조사하고, 유저 프로필과 매칭하고, "
    "면접까지 코치하는 완성형 AI 취업 코치입니다."
)
add_para(
    "개인정보를 최대한 보호하는 로컬 우선 설계, 포인트 기반의 유연한 과금, "
    "어드민에서 프롬프트를 실시간 개선하는 운영 구조는 빠른 제품 개선과 높은 마진율을 동시에 달성합니다."
)

doc.add_paragraph()
add_para(
    "한국 취업 준비생 320만 명 × 연간 12개 지원 = 연간 3,840만 건의 자소서 시장.",
    bold=True, color=BLUE
)
add_para(
    "SelfPR은 이 시장의 1%만 점유해도 연 192억 원 매출을 달성할 수 있습니다.",
    bold=True, color=BLUE
)

doc.add_paragraph()
p_foot = doc.add_paragraph()
p_foot.paragraph_format.space_before = Pt(24)
r_foot = p_foot.add_run("본 사업계획서는 2026년 3월 기준으로 작성되었습니다.")
r_foot.font.size = Pt(9)
r_foot.font.color.rgb = SLATE
p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ─── 저장 ───────────────────────────────────────────────────
out_path = "/Users/gisoojeong/Documents/aiops/selfpr/docs/SelfPR_사업계획서_v1.0.docx"
doc.save(out_path)
print(f"✅ 저장 완료: {out_path}")
